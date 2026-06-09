"""Shared fixtures: a realistic legal mailbox built fresh per test.

Two messages: the deterministic sample.eml shipped by mail2md, and a
synthetic privileged email carrying a .docx (must convert), a .csv (native
upload), a .md (skip — already Markdown), with a Windows-illegal subject and
attachment name to exercise sanitization end to end.
"""

from __future__ import annotations

import io
import shutil
from email.message import EmailMessage
from pathlib import Path

import pytest

MAIL2MD_FIXTURE = (
    Path(__file__).resolve().parents[2]
    / "mail2md-computer-use" / "tests" / "fixtures" / "sample.eml"
)


def build_privileged_eml(path: Path) -> None:
    import docx  # python-docx, installed via omniconvert's [documents] extra

    doc = docx.Document()
    doc.add_heading("Settlement Memorandum", 1)
    doc.add_paragraph("ATTORNEY-CLIENT PRIVILEGED work product — Lucerne matter.")
    buf = io.BytesIO()
    doc.save(buf)

    msg = EmailMessage()
    msg["Subject"] = "RE: Lucerne — settlement memo CON: draft? <final>"
    msg["From"] = "Paula Counsel <paula@firmlaw.com>"
    msg["To"] = "Fred M <fred@client.example>, Joe Client <joe@client.example>"
    msg["Date"] = "Tue, 14 Apr 2026 09:30:00 -0700"
    msg["Message-ID"] = "<memo-001@firmlaw.com>"
    msg.set_content(
        "Privileged and confidential — attorney work product.\n"
        "Attached: settlement memo, damages csv, notes.\nDo not forward."
    )
    msg.add_attachment(
        buf.getvalue(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Settlement Memo: draft? <v2>.docx",
    )
    msg.add_attachment(
        b"date,amount\n2026-01-05,12000\n", maintype="text", subtype="csv",
        filename="damages.csv",
    )
    msg.add_attachment(
        b"# Notes\nAlready markdown.\n", maintype="text", subtype="markdown",
        filename="notes.md",
    )
    path.write_bytes(bytes(msg))


@pytest.fixture
def workspace(tmp_path: Path) -> Path:
    """A case workspace with two mail files dropped in 01_mail_in."""
    inbox = tmp_path / "01_mail_in"
    inbox.mkdir()
    shutil.copy(MAIL2MD_FIXTURE, inbox / "sample.eml")
    build_privileged_eml(inbox / "memo.eml")
    return tmp_path


@pytest.fixture
def pipeline(workspace: Path):
    from lexintake.config import LexConfig
    from lexintake.pipeline import LexIntakePipeline

    cfg = LexConfig(notebook="TEST-NB", backend="stub", workspace=workspace)
    return LexIntakePipeline(cfg)
